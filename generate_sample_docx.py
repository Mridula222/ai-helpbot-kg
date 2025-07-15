from docx import Document

doc = Document()
doc.add_heading("Sample DOCX Document", level=1)
doc.add_paragraph("This is a test paragraph in a Word document.")
doc.save("data_ingestion/docs/sample.docx")
